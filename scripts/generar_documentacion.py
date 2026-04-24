from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(22)
    return p

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    run.font.size = Pt(16)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    run.font.size = Pt(13)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix + ': ')
        run.bold = True
    p.add_run(text)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x17, 0x17, 0x17)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.pPr.append(shading)
    return p

def add_info_box(doc, text, color='DEEAF1'):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run('   ' + text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)
    doc.add_paragraph()
    return table

def make_table(doc, headers, rows_data, header_color='1F497D', alt_color='DEEAF1'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color)
        tcPr.append(shd)
    for idx, row_data in enumerate(rows_data):
        row = table.add_row()
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
            if idx % 2 == 0:
                tc = row.cells[j]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), alt_color)
                tcPr.append(shd)
    doc.add_paragraph()
    return table

# ═════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_title(doc, 'ELEKTRA S.R.L.')
p = doc.add_paragraph('Plataforma de Datos — Documentacion del Proyecto')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(14)
p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
doc.add_paragraph()
p2 = doc.add_paragraph(f'Santa Fe, Argentina  -  {datetime.date.today().strftime("%B %Y")}')
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.runs[0].font.size = Pt(11)
p2.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph()
p3 = doc.add_paragraph('VERSION COMPLETA - Fases 1 a 5 finalizadas')
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.runs[0].font.size = Pt(10)
p3.runs[0].bold = True
p3.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# INDICE
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, 'Contenido')
items = [
    '1.   Que es este proyecto y para que sirve',
    '2.   Las herramientas que usamos',
    '3.   Estructura de carpetas y archivos',
    '4.   El archivo Excel - que habia y como estaba',
    '5.   FASE 1 - Limpieza y normalizacion de datos',
    '       5.1  Script 01 - INSUMOS',
    '       5.2  Script 02 - EQUIPOS',
    '       5.3  Script 03 - CLIENTES',
    '       5.4  Script 04 - PROVEEDORES',
    '       5.5  Script 05 - PRODUCTOS y CATEGORIAS',
    '       5.6  Script 06 - COSTOS DE PRODUCCION',
    '6.   FASE 2 - Esquema SQL y base de datos relacional',
    '       6.1  Que es una base de datos relacional',
    '       6.2  Las tablas creadas',
    '       6.3  Las vistas de negocio',
    '7.   FASE 3 - Funciones de negocio en SQL',
    '       7.1  Vistas analiticas',
    '       7.2  Funciones de consulta',
    '       7.3  Procedimientos de actualizacion de precios',
    '8.   FASE 4 - Dashboard con Streamlit',
    '       8.1  Como levantar el dashboard',
    '       8.2  Pagina Inicio',
    '       8.3  Pagina Precios',
    '       8.4  Pagina Rentabilidad',
    '       8.5  Pagina Costos',
    '9.   FASE 5 - Automatizaciones',
    '       9.1  Exportador de listas a Excel formateado',
    '       9.2  Generador de facturas A/B en PDF',
    '       9.3  API REST de precios con FastAPI',
    '10.  Resumen de resultados',
    '11.  Como usar el sistema dia a dia',
]
for item in items:
    p = doc.add_paragraph(item)
    p.style = doc.styles['Normal']
    p.paragraph_format.space_after = Pt(3)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. QUE ES ESTE PROYECTO
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '1.  Que es este proyecto y para que sirve')
add_body(doc,
    'Elektra S.R.L. acumulo mas de 20 anos de informacion operativa dentro de un unico '
    'archivo de Excel con 32 hojas. La solucion que construimos transforma todo ese Excel '
    'en una plataforma de datos moderna con cinco componentes:')
add_bullet(doc, 'Una base de datos PostgreSQL con todos los datos limpios y relacionados entre si.', 'Base de datos')
add_bullet(doc, 'Funciones SQL que calculan costos, margenes y simulan aumentos de precios.', 'Logica de negocio')
add_bullet(doc, 'Un dashboard web con graficos interactivos accesible desde cualquier navegador.', 'Dashboard')
add_bullet(doc, 'Un exportador que genera listas de precios en Excel con formato profesional.', 'Excel automatico')
add_bullet(doc, 'Un generador de facturas A y B en PDF y una API REST para consultar precios.', 'Facturas y API')
add_info_box(doc,
    'Analogia simple: si el Excel era como tener todos los documentos de la empresa '
    'en una pila de papeles, la plataforma es como tener un sistema de gestion propio '
    'con busqueda instantanea, reportes automaticos y acceso desde cualquier dispositivo.')
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 2. HERRAMIENTAS
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '2.  Las herramientas que usamos')
herramientas = [
    ('Python 3.14', 'El lenguaje de programacion con el que escribimos todos los scripts. Es el mas usado en el mundo para analisis de datos y automatizacion.'),
    ('pandas / openpyxl / xlrd', 'Librerias para leer el Excel, limpiar columnas, eliminar filas vacias y reorganizar informacion.'),
    ('SQLAlchemy', 'El puente entre Python y la base de datos PostgreSQL.'),
    ('PostgreSQL 18', 'La base de datos donde guardamos toda la informacion limpia. Gratuita, robusta, usada por empresas de todo el mundo.'),
    ('pgAdmin 4', 'Interfaz visual para explorar y consultar la base de datos.'),
    ('Streamlit', 'Framework Python para crear aplicaciones web interactivas sin necesidad de saber HTML o JavaScript.'),
    ('Plotly', 'Libreria de graficos interactivos usada dentro del dashboard.'),
    ('ReportLab', 'Libreria para generar archivos PDF desde Python (facturas).'),
    ('FastAPI + Uvicorn', 'Framework para crear APIs REST. Permite consultar precios desde cualquier aplicacion o celular.'),
    ('python-dotenv', 'Guarda la contrasena de la base de datos en un archivo separado (.env) que nunca se sube a internet.'),
]
for nombre, desc in herramientas:
    add_h2(doc, nombre)
    add_body(doc, desc)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3. ESTRUCTURA DE CARPETAS
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '3.  Estructura de carpetas y archivos')
add_code_block(doc,
    'ELEKTRA-plataforma-de-datos/\n'
    '|\n'
    '+-- ELEKTRA 2026.xls              <- Excel original (fuente)\n'
    '+-- .env                          <- Credenciales DB (nunca subir)\n'
    '+-- .gitignore                    <- Ignora .env en Git\n'
    '|\n'
    '+-- scripts/\n'
    '|   +-- 01_limpieza_insumos.py\n'
    '|   +-- 02_limpieza_equipos.py\n'
    '|   +-- 03_limpieza_clientes.py\n'
    '|   +-- 04_limpieza_proveedores.py\n'
    '|   +-- 05_limpieza_productos.py\n'
    '|   +-- 06_limpieza_costos.py\n'
    '|   +-- 07_crear_esquema_sql.py\n'
    '|   +-- 08_funciones_negocio.py\n'
    '|   +-- 09_exportar_listas_excel.py\n'
    '|   +-- 10_generar_factura_pdf.py\n'
    '|   +-- 11_api_precios.py\n'
    '|   +-- generar_documentacion.py\n'
    '|\n'
    '+-- app/                          <- Dashboard Streamlit\n'
    '|   +-- Inicio.py\n'
    '|   +-- pages/\n'
    '|   |   +-- 1_Precios.py\n'
    '|   |   +-- 2_Rentabilidad.py\n'
    '|   |   +-- 3_Clientes.py\n'
    '|   |   +-- 4_Costos.py\n'
    '|   +-- utils/\n'
    '|       +-- db.py\n'
    '|\n'
    '+-- data/                         <- CSVs limpios + Excel exportado\n'
    '+-- facturas/                     <- PDFs generados'
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 4. EL EXCEL ORIGINAL
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '4.  El archivo Excel - que habia y como estaba')
add_body(doc, 'El archivo ELEKTRA 2026.xls tenia 32 hojas. Procesamos las hojas estructuradas:')
make_table(doc,
    ['Hoja', 'Contenido', 'Filas'],
    [
        ('INSUMOS',      'Materiales: codigo, descripcion, precio ARS/USD, proveedor, rubro, fecha', '3.358'),
        ('EQUIPOS',      'Productos terminados con precios publico y distribuidor en ARS y USD', '2.132'),
        ('CLIENTES',     'Clientes con CUIT, provincia, tipo IVA (172 columnas en total!)', '1.763'),
        ('PROVEEDORES',  'Proveedores con telefono, rubro, email, cotizacion dolar', '259'),
        ('PRODUCTOS',    'Catalogo con precios, IVA y categorias', '2.131'),
        ('CATEGORIAS',   'Tabla maestra de familias de productos', '19'),
        ('COSTOS',       'Costos de produccion por equipo (estructura en bloques)', '12.207'),
    ]
)
add_body(doc, 'Principales problemas encontrados en el Excel original:')
add_bullet(doc, 'Nombres de columnas con acentos, puntos y espacios que Python no puede manejar directamente.', 'Encoding roto')
add_bullet(doc, 'CLIENTES tenia 172 columnas, la mayoria sin nombre ni datos utiles.', 'Columnas vacias')
add_bullet(doc, 'Columnas numericas con texto, fechas guardadas como texto.', 'Tipos de datos mixtos')
add_bullet(doc, 'El Excel usaba filas en blanco como separadores visuales.', 'Filas vacias')
add_bullet(doc, 'COSTOS no era una tabla plana sino bloques: nombre de equipo, lista de materiales, siguiente equipo...', 'Estructura compleja')
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 5. FASE 1
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '5.  FASE 1 - Limpieza y normalizacion de datos')
add_body(doc,
    'La limpieza de datos es el paso mas importante. Cada script de limpieza sigue el mismo proceso: '
    'leer la hoja del Excel, renombrar columnas a snake_case, eliminar filas vacias, '
    'convertir tipos de datos, normalizar texto, eliminar duplicados, guardar CSV en /data y cargar a PostgreSQL.')

resultados_f1 = [
    ('01_limpieza_insumos.py',    'INSUMOS',     '3.347', 'Columna extra con numero como nombre eliminada. Encoding de "Dolar" corregido.'),
    ('02_limpieza_equipos.py',    'EQUIPOS',     '2.131', 'Columna ARTICULO duplicada eliminada. 2 columnas Unnamed eliminadas.'),
    ('03_limpieza_clientes.py',   'CLIENTES',    '1.394', '161 columnas vacias descartadas. Telefono armado de area + numero.'),
    ('04_limpieza_proveedores.py','PROVEEDORES',  '254', 'Encoding de "dolar" corregido. Fechas de formato mixto normalizadas.'),
    ('05_limpieza_productos.py',  'PRODUCTOS',   '2.131', 'Encoding de "Dolares" e "IdCategoria" corregido.'),
    ('05_limpieza_productos.py',  'CATEGORIAS',     '19', 'Tabla maestra cargada primero como referencia.'),
    ('06_limpieza_costos.py',     'COSTOS',      '5.752', 'Bloques de equipo parseados con algoritmo fila por fila. 2.072 equipos distintos.'),
]
make_table(doc,
    ['Script', 'Hoja', 'Filas limpias', 'Principales ajustes'],
    resultados_f1
)

add_h2(doc, '5.3  CLIENTES - el caso mas complejo')
add_body(doc,
    'La hoja de CLIENTES merece mencion especial. Tenia 172 columnas, la gran mayoria '
    'sin nombre ni datos. Identificamos las columnas utiles por posicion: '
    'col 1=codigo, 2=nombre, 3=fantasia, 4=contacto, 5=direccion, 6=codigo postal, '
    '7=ciudad, 9-10=telefono, 30-31=telefono alternativo, 39=CUIT, 170=tipo IVA, 171=provincia. '
    'Descartamos las 161 columnas restantes.')

add_h2(doc, '5.6  COSTOS - estructura en bloques')
add_body(doc,
    'Los costos estaban organizados en bloques: una fila con el nombre del equipo, '
    'luego filas de materiales, luego el siguiente equipo. Escribimos un algoritmo '
    'que recorre fila por fila y detecta cuando empieza un nuevo equipo (fila con nombre '
    'pero sin codigo de material) para transformar los bloques en una tabla plana.')
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 6. FASE 2
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '6.  FASE 2 - Esquema SQL y base de datos relacional')

add_h2(doc, '6.1  Que es una base de datos relacional')
add_body(doc,
    'Una base de datos relacional conecta las tablas entre si a traves de claves. '
    'Por ejemplo, en el Excel si un insumo tiene "proveedor_codigo = 22", hay que ir '
    'a buscar manualmente en la hoja de proveedores que empresa es. En la base relacional '
    'esa conexion es automatica y permite hacer consultas que cruzan informacion de '
    'multiples tablas en milisegundos.')

add_h2(doc, '6.2  Las tablas creadas (script 07)')
make_table(doc,
    ['Tabla', 'Que contiene', 'Relacion'],
    [
        ('categorias',     '19 familias de productos',                  'Tabla maestra'),
        ('tipos_iva',      'Tipos de IVA posibles',                     'Tabla maestra'),
        ('provincias',     'Provincias argentinas',                     'Tabla maestra'),
        ('proveedores',    '254 proveedores con todos sus datos',       'Independiente'),
        ('clientes',       '1.394 clientes con CUIT y tipo IVA',        'Independiente'),
        ('insumos',        '3.347 materiales con precios ARS y USD',    'FK -> proveedores'),
        ('productos',      '2.131 productos del catalogo de venta',     'FK -> categorias'),
        ('equipos',        '2.131 equipos fabricados',                  'Independiente'),
        ('costos_equipos', '5.752 registros material por equipo',       'Por nombre de equipo'),
    ]
)

add_h2(doc, '6.3  Las vistas del script 07')
add_body(doc, 'Las vistas son consultas guardadas que se actualizan solas con los datos mas recientes:')
add_bullet(doc, 'Costo total de produccion de cada equipo sumando todos sus materiales.', 'v_costo_por_equipo')
add_bullet(doc, 'Margen real, precio publico y precio distribuidor por equipo.', 'v_margen_equipos')
add_bullet(doc, 'Lista de insumos con stock = 0 o sin dato.', 'v_insumos_sin_stock')
add_bullet(doc, 'Clientes agrupados por provincia y tipo de IVA.', 'v_clientes_por_provincia')
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 7. FASE 3
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '7.  FASE 3 - Funciones de negocio en SQL (script 08)')
add_body(doc,
    'La Fase 3 agrego logica de negocio directamente en la base de datos: '
    'vistas analiticas mas avanzadas, funciones de consulta y procedimientos '
    'que modifican precios. Todo vive en el script 08_funciones_negocio.py.')

add_h2(doc, '7.1  Vistas analiticas nuevas')
make_table(doc,
    ['Vista', 'Para que sirve'],
    [
        ('v_rentabilidad_por_categoria', 'Precio promedio, precio USD y margen al distribuidor por cada familia de productos.'),
        ('v_lista_precios_publica',      'Lista de precios al publico lista para exportar: sin IVA, con IVA, en USD.'),
        ('v_lista_precios_distribuidor', 'Lista de precios para distribuidores con descuento.'),
        ('v_materiales_mas_usados',      'Materiales que aparecen en mas equipos distintos con impacto en costos totales.'),
        ('v_equipos_rentabilidad',       'Margen real, margen al distribuidor y precio USD por cada equipo fabricado.'),
        ('v_insumos_caros_sin_usd',      'Insumos sin precio en dolares, con estimacion calculada desde ARS / cotizacion.'),
    ]
)

add_h2(doc, '7.2  Funciones de consulta')
add_body(doc, 'Las funciones se llaman con SELECT y devuelven datos sin modificar nada:')
add_code_block(doc,
    '-- Ver el desglose completo de materiales del equipo VIBRO:\n'
    'SELECT * FROM fn_detalle_costo_equipo(\'VIBRO\');\n'
    '\n'
    '-- Ver el costo total de produccion del RADAR:\n'
    'SELECT fn_costo_total_equipo(\'RADAR\');\n'
    '\n'
    '-- Convertir 100.000 pesos a dolares con cotizacion 1.500:\n'
    'SELECT fn_ars_a_usd(100000, 1500);\n'
    '\n'
    '-- Simular como quedarian los precios con un aumento del 20%%:\n'
    'SELECT * FROM fn_simular_aumento_insumos(20);\n'
    'SELECT * FROM fn_simular_aumento_equipos(20);'
)

add_h2(doc, '7.3  Procedimientos de actualizacion de precios')
add_body(doc,
    'Los procedimientos modifican los datos de la base. Se llaman con CALL y tienen '
    'una validacion que impide porcentajes invalidos (menores a 0 o mayores a 500):')
add_code_block(doc,
    '-- Aplicar 15.5%% de aumento solo a insumos:\n'
    'CALL sp_aumento_insumos(15.5);\n'
    '\n'
    '-- Aplicar 15.5%% de aumento solo a equipos:\n'
    'CALL sp_aumento_equipos(15.5);\n'
    '\n'
    '-- Aplicar 15.5%% de aumento solo a productos:\n'
    'CALL sp_aumento_productos(15.5);\n'
    '\n'
    '-- Aplicar 15.5%% a todo de una vez (insumos + equipos + productos):\n'
    'CALL sp_aumento_general(15.5);'
)
add_info_box(doc,
    'IMPORTANTE: Siempre usar primero fn_simular_aumento_insumos() para ver como '
    'quedarian los precios ANTES de llamar al procedimiento que los modifica. '
    'Una vez aplicado el aumento, la base queda actualizada.', 'FCE4D6')
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 8. FASE 4
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '8.  FASE 4 - Dashboard con Streamlit')

add_h2(doc, '8.1  Como levantar el dashboard')
add_body(doc, 'Desde la terminal, pararse en la carpeta app/ y ejecutar:')
add_code_block(doc, 'py -m streamlit run Inicio.py')
add_body(doc, 'Se abre automaticamente en el navegador en http://localhost:8501')

add_h2(doc, '8.2  Pagina Inicio')
add_body(doc,
    'La pagina principal muestra un resumen ejecutivo del negocio: '
    '5 contadores (insumos, equipos, clientes, proveedores, equipos con costo calculado), '
    'top 5 categorias por precio promedio, top 5 equipos por margen real, '
    'grafico de clientes por provincia y materiales mas usados en produccion.')

add_h2(doc, '8.3  Pagina Precios')
add_body(doc,
    'Lista de precios interactiva con filtros por categoria y buscador por nombre. '
    'Permite elegir entre lista publica y lista distribuidor. '
    'Tiene un control deslizante para ajustar la cotizacion del dolar y ver '
    'al instante cuanto vale cada producto en USD. '
    'Incluye boton para descargar la lista filtrada como archivo Excel.')

add_h2(doc, '8.4  Pagina Rentabilidad')
add_body(doc, 'Tres pestanas de analisis:')
add_bullet(doc, 'Grafico de barras de precio promedio por categoria y margen al distribuidor por categoria.', 'Por categoria')
add_bullet(doc, 'Grafico scatter (nube de puntos) de costo vs precio con linea de break-even. Un punto encima de la linea roja significa ganancia; debajo significa perdida.', 'Por equipo')
add_bullet(doc, 'Selector de equipo especifico que muestra el grafico de torta con la composicion del costo (que porcentaje representa cada material) y la tabla de materiales completa.', 'Desglose de costo')

add_h2(doc, '8.5  Pagina Costos')
add_body(doc, 'Tres pestanas:')
add_bullet(doc, 'Buscador de insumos por descripcion y rubro, con grafico de precio promedio por rubro.', 'Insumos')
add_bullet(doc, 'Ranking de equipos por costo de produccion total con slider para elegir cuantos mostrar, y desglose detallado del equipo seleccionado.', 'Equipos')
add_bullet(doc, 'Simulador (muestra como quedarian los precios sin modificarlos) y aplicador de aumentos con checkbox de confirmacion para evitar errores accidentales.', 'Actualizar precios')
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 9. FASE 5
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '9.  FASE 5 - Automatizaciones')

add_h2(doc, '9.1  Exportador de listas a Excel formateado (script 09)')
add_body(doc,
    'El script 09_exportar_listas_excel.py genera un archivo Excel profesional con tres hojas, '
    'cada una con su propio esquema de colores, filas alternadas, separadores por '
    'categoria o rubro, columnas con ancho automatico y encabezados fijos (freeze panes):')
make_table(doc,
    ['Hoja', 'Color', 'Contenido'],
    [
        ('Lista Publica',      'Azul',    '1.980 productos con precio sin IVA, con IVA, en USD y descuento distribuidor'),
        ('Lista Distribuidor', 'Verde',   '1.979 productos con precio distribuidor sin y con IVA'),
        ('Insumos',            'Naranja', '3.342 materiales con precio ARS, ARS+IVA, USD, cotizacion y stock'),
    ]
)
add_body(doc, 'Para ejecutarlo:')
add_code_block(doc, 'py scripts/09_exportar_listas_excel.py')
add_body(doc, 'El archivo se guarda en data/ELEKTRA_Listas_Precios_AAAAMMDD.xlsx con la fecha del dia.')

add_h2(doc, '9.2  Generador de facturas A/B en PDF (script 10)')
add_body(doc,
    'El script 10_generar_factura_pdf.py genera una factura en PDF lista para imprimir. '
    'El tipo de factura (A o B) se determina automaticamente segun el tipo de IVA del cliente:')
add_bullet(doc, 'Factura A: para clientes "Responsable Inscripto". Muestra IVA discriminado.', 'Factura A')
add_bullet(doc, 'Factura B: para "Consumidor Final", "Monotributista" y "Exento". IVA incluido en el precio.', 'Factura B')
add_body(doc, 'Contenido de la factura generada:')
add_bullet(doc, 'Encabezado con datos de Elektra SRL: nombre, direccion, CUIT, inicio de actividades.')
add_bullet(doc, 'Recuadro con la letra A o B bien visible (requerimiento AFIP).')
add_bullet(doc, 'Datos del cliente: razon social, direccion, ciudad, CUIT, tipo de IVA.')
add_bullet(doc, 'Tabla de items: descripcion, cantidad, precio unitario, subtotal y (si es tipo A) IVA.')
add_bullet(doc, 'Totales: subtotal neto, IVA y total a pagar.')
add_body(doc, 'Formas de usar el script:')
add_code_block(doc,
    '# Generar factura de demostracion (toma el primer cliente con CUIT):\n'
    'py scripts/10_generar_factura_pdf.py --demo\n'
    '\n'
    '# Generar factura para el cliente 15 con 1 RADAR y 2 VIBROs:\n'
    'py scripts/10_generar_factura_pdf.py --cliente 15 --items "RADAR:1,VIBRO:2"\n'
    '\n'
    '# Con numero de factura personalizado:\n'
    'py scripts/10_generar_factura_pdf.py --cliente 15 --items "RADAR:1" --numero "0001-00000042"'
)
add_body(doc, 'Los PDFs se guardan en la carpeta facturas/ con el nombre: FACTURA_A_NUMERO_CLIENTE.pdf')

add_h2(doc, '9.3  API REST de precios con FastAPI (script 11)')
add_body(doc,
    'El script 11_api_precios.py crea una API REST que permite consultar precios, '
    'insumos y costos desde cualquier aplicacion, celular o sistema externo '
    'mediante peticiones HTTP simples.')
add_body(doc, 'Como levantarla:')
add_code_block(doc, 'py -m uvicorn scripts/11_api_precios:app --reload --port 8000')
add_body(doc,
    'Una vez corriendo, FastAPI genera automaticamente una documentacion interactiva '
    'en http://localhost:8000/docs donde se puede probar cada endpoint con un formulario visual.')
add_body(doc, 'Endpoints disponibles:')
make_table(doc,
    ['Endpoint', 'Metodo', 'Que devuelve'],
    [
        ('/productos',                  'GET', 'Lista de precios al publico con filtros por categoria y nombre'),
        ('/productos/{nombre}',         'GET', 'Precio publico Y distribuidor de un producto especifico'),
        ('/insumos',                    'GET', 'Lista de insumos con filtros por rubro y descripcion'),
        ('/insumos/{codigo}',           'GET', 'Detalle completo de un insumo por codigo'),
        ('/equipos',                    'GET', 'Lista de equipos con margen real y precios'),
        ('/equipos/{nombre}/costo',     'GET', 'Costo de produccion completo de un equipo con desglose'),
        ('/clientes',                   'GET', 'Lista de clientes con filtros por provincia y tipo IVA'),
        ('/clientes/{codigo}',          'GET', 'Datos completos de un cliente'),
        ('/categorias',                 'GET', 'Categorias con precio promedio y margen'),
        ('/cotizacion',                 'GET', 'Cotizaciones de dolar registradas por proveedor'),
        ('/stock-critico',              'GET', 'Insumos con stock = 0 o sin dato'),
    ]
)
add_body(doc, 'Ejemplo de consulta desde el navegador o cualquier programa:')
add_code_block(doc,
    'http://localhost:8000/productos/RADAR\n'
    '  -> Devuelve: precio sin IVA, con IVA, en USD, precio distribuidor\n'
    '\n'
    'http://localhost:8000/equipos/VIBRO/costo\n'
    '  -> Devuelve: costo total y lista de todos los materiales con cantidades\n'
    '\n'
    'http://localhost:8000/clientes?provincia=SANTA FE&tipo_iva=responsable\n'
    '  -> Devuelve: lista de clientes RI de Santa Fe'
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 10. RESUMEN DE RESULTADOS
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '10.  Resumen de resultados')
add_body(doc, 'Al completar las 5 fases, la plataforma contiene:')
make_table(doc,
    ['Componente', 'Detalle', 'Estado'],
    [
        ('Base de datos PostgreSQL',  'elektra_srl con 9 tablas y datos limpios', 'ACTIVO'),
        ('Insumos cargados',          '3.347 materiales con precios ARS y USD', 'ACTIVO'),
        ('Equipos cargados',          '2.131 equipos con precios y margenes', 'ACTIVO'),
        ('Clientes cargados',         '1.394 clientes con CUIT y tipo IVA', 'ACTIVO'),
        ('Proveedores cargados',      '254 proveedores con cotizacion dolar', 'ACTIVO'),
        ('Costos de produccion',      '5.752 registros - 2.072 equipos distintos', 'ACTIVO'),
        ('Vistas SQL',                '10 vistas analiticas de negocio', 'ACTIVO'),
        ('Funciones SQL',             '7 funciones de consulta y conversion', 'ACTIVO'),
        ('Procedimientos SQL',        '4 procedimientos de actualizacion de precios', 'ACTIVO'),
        ('Dashboard Streamlit',       '4 paginas: Inicio, Precios, Rentabilidad, Costos', 'ACTIVO'),
        ('Exportador Excel',          'Genera lista publica + distribuidor + insumos formateados', 'ACTIVO'),
        ('Generador PDF',             'Facturas A y B con datos del cliente y detalle de items', 'ACTIVO'),
        ('API REST',                  '11 endpoints documentados con FastAPI', 'ACTIVO'),
        ('Documentacion',             'Este documento Word generado automaticamente', 'ACTIVO'),
    ]
)
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 11. COMO USAR EL SISTEMA DIA A DIA
# ═════════════════════════════════════════════════════════════════════════════
add_h1(doc, '11.  Como usar el sistema dia a dia')

add_h2(doc, 'Consultar precios')
add_body(doc, 'Opcion 1 - Dashboard (la mas comoda):')
add_code_block(doc, 'cd app\npy -m streamlit run Inicio.py')
add_body(doc, 'Opcion 2 - API desde el navegador:')
add_code_block(doc,
    'py -m uvicorn scripts/11_api_precios:app --port 8000\n'
    '# Abrir: http://localhost:8000/productos/RADAR'
)

add_h2(doc, 'Aplicar un aumento de precios')
add_body(doc, 'Desde pgAdmin (recomendado para mayor control):')
add_code_block(doc,
    '-- Primero SIMULAR (no modifica nada):\n'
    'SELECT * FROM fn_simular_aumento_insumos(15);\n'
    '\n'
    '-- Si el resultado es correcto, APLICAR:\n'
    'CALL sp_aumento_general(15);'
)
add_body(doc, 'Desde el dashboard (mas sencillo):')
add_bullet(doc, 'Ir a la pagina "Costos".')
add_bullet(doc, 'Abrir la pestana "Actualizar precios".')
add_bullet(doc, 'Ingresar el porcentaje en "Simular" y verificar los precios nuevos.')
add_bullet(doc, 'Si esta correcto, ingresar el mismo porcentaje en "Aplicar", tildar la confirmacion y presionar el boton.')

add_h2(doc, 'Generar una lista de precios en Excel')
add_code_block(doc, 'py scripts/09_exportar_listas_excel.py')
add_body(doc, 'El archivo queda en data/ELEKTRA_Listas_Precios_AAAAMMDD.xlsx')

add_h2(doc, 'Generar una factura PDF')
add_code_block(doc,
    '# Reemplazar 15 por el codigo del cliente y los productos reales:\n'
    'py scripts/10_generar_factura_pdf.py --cliente 15 --items "RADAR:1,VIBRO:2"'
)
add_body(doc, 'La factura queda en la carpeta facturas/')

add_h2(doc, 'Actualizar los datos cuando cambia el Excel')
add_body(doc, 'Si Elektra actualiza el archivo Excel fuente, ejecutar los scripts de limpieza en orden:')
add_code_block(doc,
    'cd scripts\n'
    'py 01_limpieza_insumos.py\n'
    'py 02_limpieza_equipos.py\n'
    'py 03_limpieza_clientes.py\n'
    'py 04_limpieza_proveedores.py\n'
    'py 05_limpieza_productos.py\n'
    'py 06_limpieza_costos.py'
)
add_body(doc,
    'Los scripts usan if_exists="replace" lo que significa que borran la tabla vieja '
    'y cargan los datos nuevos. El dashboard y la API reflejan los cambios automaticamente.')

# ── Pie de pagina ─────────────────────────────────────────────────────────────
section = doc.sections[0]
footer = section.footer
p_footer = footer.paragraphs[0]
p_footer.text = f'Elektra S.R.L. - Plataforma de Datos  -  Generado el {datetime.date.today().strftime("%d/%m/%Y")}  -  Version completa (Fases 1-5)'
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.runs[0].font.size = Pt(8)
p_footer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Guardar ───────────────────────────────────────────────────────────────────
output = '../ELEKTRA_Documentacion_Proyecto.docx'
doc.save(output)
print(f'Documento guardado: {output}')
